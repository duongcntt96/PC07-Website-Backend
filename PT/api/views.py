from rest_framework import viewsets
#  For custom allow API
from rest_framework import mixins
# Import Serializer
from .serializers import ImageSerializer, PhuongtienSerializer, TeamSerializer, ChungloaiSerializer, DinhmucnhienlieuSerializer
from .serializers import NhattrinhxeSerializer, TrangthaiSerializer, ChatluongSerializer
# Import Models
from PT.models import Image, Phuong_tien, Chung_loai, Dinh_muc_nhien_lieu, Nhat_trinh_xe, Trang_thai, Chat_luong
# Import permissions
from rest_framework import permissions
from api.permissions import IsPostOrAdmin, IsAuthenticatedOrReadOnly, ObjectPermission

from api.paginations import CustomPagination

from rest_framework.filters import SearchFilter, OrderingFilter
from django_filters.rest_framework import DjangoFilterBackend

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status

from django.contrib.auth.decorators import login_required
from api.decorators import role_required

from profiles.models import Team


class PhuongtienViewSet(viewsets.ModelViewSet):
    queryset = Phuong_tien.objects.all()
    serializer_class = PhuongtienSerializer
    pagination_class = CustomPagination
    # permission_classes = [ObjectPermission]

    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['noi_bo_tri', 'chung_loai',
                        'quan_ly', 'trang_thai', 'chat_luong']
    search_fields = ["ten", "chung_loai__ten"]

    def create(self, request):
        serializer = PhuongtienSerializer(data=request.data)
        if serializer.is_valid():
            if Phuong_tien.objects.get(pk=request.data.get('noi_bo_tri')).quan_ly in request.user.profile.all()[0].team.all():
                serializer.save()
                import docx

                def doc_replace(doc, findText, replaceText):
                    for p in doc.paragraphs:
                        print(p.text)
                        if findText in p.text:
                            inline = p.runs
                            # Loop added to work with runs (strings with same style)
                            for i in range(len(inline)):
                                if findText in inline[i].text:
                                    text = inline[i].text.replace(
                                        findText, replaceText)
                                    inline[i].text = text

                data = []
                doc = docx.Document("E:/6. QUYẾT ĐỊNH CẤP GIẤY.docx")
                findTexts = ["<tên cơ sở>", "<địa chỉ>", "<số điện thoại>",
                             "<cơ quan quản lý>", "<số điện thoại cơ quan quản lý>"]
                for findText in findTexts:
                    replace = input("Nhập "+findText+":")
                    data.append(replace)
                i = 0
                for findText in findTexts:
                    print(findText+"=>"+data[i])
                    doc_replace(doc, findText, data[i])
                    i += 1
                doc.save('E:/Bìa.docx')
                print("OK")

                # doc2 = docx.Document("E:/Phiếu chiến thuật.docx")
                # i = 0
                # for findText in findTexts:
                #     print(findText+"=>"+data[i])
                #     doc_replace(doc2, findText, data[i])
                #     i += 1
                # doc2.save('Phiếu chiến thuật.docx')
                # print("OK")
                return Response({'message': 'Post saved !'}, status=status.HTTP_200_OK)
            else:
                return Response({'message': 'Not permission', }, status=status.HTTP_405_METHOD_NOT_ALLOWED)
        else:
            return Response({'message': 'Not found', }, status=status.HTTP_404_NOT_FOUND)


class PhuongtienOfToAPIView(APIView):
    # permission_classes = [permissions.IsAuthenticated]
    # @login_required
    @role_required(['Admin'])
    def get(self, request):
        teamID = request.query_params.get('teamID')
        print(teamID)
        queryset = Phuong_tien.objects.filter(nguoi_quan_ly=teamID)

        rp = PhuongtienSerializer(queryset, many=True)
        return Response(rp.data, status=status.HTTP_200_OK)

    def post(self, request):
        from .serializers import UserSerializer
        serializer = UserSerializer(data=request.data)
        if (serializer.is_valid()):
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class TeamViewSet(viewsets.ModelViewSet):
    queryset = Team.objects.all()
    serializer_class = TeamSerializer
    pagination_class = CustomPagination


class TrangthaiViewSet(viewsets.ModelViewSet):
    queryset = Trang_thai.objects.all()
    serializer_class = TrangthaiSerializer
    pagination_class = CustomPagination
    permission_classes = [ObjectPermission]


class Chungloai_ViewSet(viewsets.ModelViewSet):
    queryset = Chung_loai.objects.exclude(stt__isnull=True)
    serializer_class = ChungloaiSerializer
    pagination_class = CustomPagination


class ChatluongViewSet(viewsets.ModelViewSet):
    queryset = Chat_luong.objects.all()
    serializer_class = ChatluongSerializer
    pagination_class = CustomPagination


class DinhmucnhienlieuViewSet(viewsets.ModelViewSet):
    queryset = Dinh_muc_nhien_lieu.objects.all()
    serializer_class = DinhmucnhienlieuSerializer
    pagination_class = CustomPagination


class NhattrinhxeViewSet(viewsets.ModelViewSet):
    queryset = Nhat_trinh_xe.objects.all()
    serializer_class = NhattrinhxeSerializer
    pagination_class = CustomPagination


class ImageViewSet(viewsets.ModelViewSet):
    queryset = Image.objects.all()
    serializer_class = ImageSerializer
    pagination_class = CustomPagination
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['phuong_tien']
    search_fields = ["phuong_tien"]
